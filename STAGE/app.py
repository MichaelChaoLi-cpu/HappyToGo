#!/usr/bin/env python3
"""
STAGE/app.py  —  HappyToGo document generation web UI
Run: python STAGE/app.py
"""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path
from flask import Flask, render_template, jsonify, request

PROJECT_ROOT = Path(__file__).parent.parent
AGENT_DIR    = PROJECT_ROOT / "AGENT"
TEMP_DIR     = PROJECT_ROOT / "_temp"
INPUT_FILE   = PROJECT_ROOT / ".input"

PYTHON = sys.executable   # same interpreter that's running this app

app = Flask(__name__)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

MD_FILES = {
    "title":        "Title.md",
    "cover_letter": "CoverLetter.md",
    "highlights":   "Highlights.md",
}

SCRIPTS = {
    "title":        "title_creator.py",
    "cover_letter": "cover_letter_creator.py",
    "highlights":   "highlight_creator.py",
    "credit":       "credit_author_statement_creator.py",
}

def run_script(script_name: str, extra_args: list[str] = []) -> dict:
    script = AGENT_DIR / script_name
    env_file = PROJECT_ROOT / ".env"

    # Build command: source .env is shell-specific; instead load vars manually
    import os
    env = os.environ.copy()
    if env_file.exists():
        for line in env_file.read_text().splitlines():
            line = line.strip()
            if line.startswith("export "):
                line = line[7:]
            if "=" in line and not line.startswith("#") and not line.startswith("source"):
                k, v = line.split("=", 1)
                env[k.strip()] = v.strip().strip("'\"")

    result = subprocess.run(
        [PYTHON, str(script)] + extra_args,
        capture_output=True, text=True, env=env, cwd=str(PROJECT_ROOT)
    )
    output = result.stdout + result.stderr
    return {"ok": result.returncode == 0, "output": output}


def get_article_link() -> str:
    import re
    if not INPUT_FILE.exists():
        return ""
    for line in INPUT_FILE.read_text().splitlines():
        m = re.match(r"""article_link\s*=\s*['"]([^'"]+)['"]""", line.strip())
        if m:
            return m.group(1)
    return ""


# ---------------------------------------------------------------------------
# Routes — pages
# ---------------------------------------------------------------------------

@app.route("/")
def index():
    return render_template("index.html")


# ---------------------------------------------------------------------------
# Routes — config (.input)  — structured fields
# ---------------------------------------------------------------------------

FIELD_DEFS = [
    {"key": "article_link",       "label": "Article Folder",       "type": "text",   "required": True},
    {"key": "file_name",          "label": "Manuscript File",      "type": "text",   "required": True},
    {"key": "journal_name",       "label": "Target Journal",       "type": "text",   "required": False},
    {"key": "authors",            "label": "Authors",              "type": "list",   "required": False},
    {"key": "corresponding",      "label": "Corresponding Author", "type": "text",   "required": False},
    {"key": "cover_letter_name",  "label": "Cover Letter Filename","type": "text",   "required": False},
    {"key": "title_name",         "label": "Title Filename",       "type": "text",   "required": False},
    {"key": "highlight_name",     "label": "Highlights Filename",  "type": "text",   "required": False},
]


def parse_input_fields(path: Path) -> dict:
    import re
    result = {}
    if not path.exists():
        return result
    content = path.read_text()
    # list fields: key={...} or key=[...]
    for m in re.finditer(r'(\w+)\s*=\s*[{\[](.*?)[}\]]', content, re.DOTALL):
        items = [s.strip().strip("'\"") for s in m.group(2).split(',') if s.strip().strip("'\"")]
        result[m.group(1)] = items
    # scalar fields
    for line in content.splitlines():
        m = re.match(r"""(\w+)\s*=\s*['"]([^'"]+)['"]""", line.strip())
        if m and m.group(1) not in result:
            result[m.group(1)] = m.group(2)
    return result


def write_input_fields(fields: dict) -> str:
    lines = []
    for fd in FIELD_DEFS:
        key = fd["key"]
        val = fields.get(key)
        if val is None or val == "" or val == []:
            continue
        if isinstance(val, list):
            items = ", ".join(f"'{v}'" for v in val if v.strip())
            lines.append(f"{key}={{{items}}}")
        else:
            lines.append(f"{key}='{val}'")
    return "\n".join(lines) + "\n"


@app.route("/api/config/fields", methods=["GET"])
def get_config_fields():
    data = parse_input_fields(INPUT_FILE)
    return jsonify({"fields": data, "defs": FIELD_DEFS})


@app.route("/api/config/fields", methods=["POST"])
def save_config_fields():
    fields = request.json.get("fields", {})
    INPUT_FILE.write_text(write_input_fields(fields))
    return jsonify({"ok": True})


# ---------------------------------------------------------------------------
# Routes — markdown files
# ---------------------------------------------------------------------------

@app.route("/api/md/<doc>", methods=["GET"])
def get_md(doc):
    fname = MD_FILES.get(doc)
    if not fname:
        return jsonify({"content": "", "exists": False})
    path = TEMP_DIR / fname
    if path.exists():
        return jsonify({"content": path.read_text(encoding="utf-8"), "exists": True})
    return jsonify({"content": "", "exists": False})


@app.route("/api/md/<doc>", methods=["POST"])
def save_md(doc):
    fname = MD_FILES.get(doc)
    if not fname:
        return jsonify({"ok": False, "error": "unknown doc"})
    TEMP_DIR.mkdir(exist_ok=True)
    path = TEMP_DIR / fname
    path.write_text(request.json.get("content", ""), encoding="utf-8")
    return jsonify({"ok": True})


# ---------------------------------------------------------------------------
# Routes — script execution
# ---------------------------------------------------------------------------

@app.route("/api/generate/<doc>", methods=["POST"])
def generate(doc):
    script = SCRIPTS.get(doc)
    if not script:
        return jsonify({"ok": False, "output": "Unknown document type"})
    result = run_script(script)
    return jsonify(result)


@app.route("/api/build/<doc>", methods=["POST"])
def build(doc):
    script = SCRIPTS.get(doc)
    if not script:
        return jsonify({"ok": False, "output": "Unknown document type"})
    result = run_script(script, ["--build"])
    return jsonify(result)


def _build_declaration_docx(dst: Path):
    """Build DeclarationStatement.docx from scratch (no template needed)."""
    from docx import Document
    from docx.oxml import OxmlElement

    _EMU_INCH = 914400
    _EMU_PT   = 12700

    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = int(8.5  * _EMU_INCH)
    sec.page_height   = int(11.0 * _EMU_INCH)
    sec.left_margin   = sec.right_margin  = _EMU_INCH        # 1"
    sec.top_margin    = sec.bottom_margin = _EMU_INCH        # 1"

    # Normal style: 10pt space after
    doc.styles["Normal"].paragraph_format.space_after = int(10 * _EMU_PT)

    # Remove default empty paragraph
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    para = doc.add_paragraph()
    para.paragraph_format.space_before = None
    para.paragraph_format.space_after  = None

    def _r(text, bold=False):
        el = OxmlElement("w:r")
        if bold:
            rPr = OxmlElement("w:rPr")
            rPr.append(OxmlElement("w:b"))
            el.append(rPr)
        t = OxmlElement("w:t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = text
        el.append(t)
        para._element.append(el)

    def _br():
        el = OxmlElement("w:r")
        el.append(OxmlElement("w:br"))
        para._element.append(el)

    _r("Declaration of interests", bold=True)
    _br()
    _r(" ")   # NBSP blank line
    _br()
    _r("☒ The authors declare that they have no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper.")
    _br()
    _r(" ")   # NBSP blank line
    _br()
    _r("☐ The authors declare the following financial interests/personal relationships which may be considered as potential competing interests:")
    _br()

    doc.save(str(dst))


@app.route("/api/copy/declaration", methods=["POST"])
def copy_declaration():
    article_link = get_article_link()
    if not article_link:
        return jsonify({"ok": False, "output": "article_link not set in .input"})
    dst = Path(article_link) / "DeclarationStatement.docx"
    try:
        _build_declaration_docx(dst)
        return jsonify({"ok": True, "output": f"DeclarationStatement saved: {dst}"})
    except Exception as e:
        return jsonify({"ok": False, "output": str(e)})


# ---------------------------------------------------------------------------
# Routes — InfoCenter / NameList
# ---------------------------------------------------------------------------

INFOCENTER_DIR  = PROJECT_ROOT / "INFOCENTER"
NAMELIST_FILE   = INFOCENTER_DIR / "NameList.json"


def _norm_person(p) -> dict:
    """Normalise a NameList entry to {name, institution, email, address, position, phone}."""
    if isinstance(p, str):
        return {"name": p.strip(), "institution": "", "email": "", "address": "", "position": "", "phone": ""}
    return {"name":        p.get("name", "").strip(),
            "institution": p.get("institution", "").strip(),
            "email":       p.get("email", "").strip(),
            "address":     p.get("address", "").strip(),
            "position":    p.get("position", "").strip(),
            "phone":       p.get("phone", "").strip()}


def load_namelist() -> list[dict]:
    import json
    if not NAMELIST_FILE.exists():
        return []
    try:
        raw = json.loads(NAMELIST_FILE.read_text(encoding="utf-8"))
        return [_norm_person(p) for p in raw]
    except Exception:
        return []


def save_namelist(persons: list[dict]):
    import json
    INFOCENTER_DIR.mkdir(exist_ok=True)
    seen, result = set(), []
    for p in persons:
        p = _norm_person(p)
        if p["name"] and p["name"] not in seen:
            seen.add(p["name"])
            result.append(p)
    NAMELIST_FILE.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")


@app.route("/api/namelist", methods=["GET"])
def get_namelist():
    return jsonify({"persons": load_namelist()})


@app.route("/api/namelist", methods=["POST"])
def update_namelist():
    """Upsert persons: update existing entry's fields if name matches, else add new."""
    new_persons = request.json.get("persons", [])
    current = {p["name"]: p for p in load_namelist()}
    for p in new_persons:
        p = _norm_person(p)
        name = p["name"]
        if not name:
            continue
        if name in current:
            for field in ("institution", "email", "address", "position", "phone"):
                if p[field]:
                    current[name][field] = p[field]
        else:
            current[name] = p
    save_namelist(list(current.values()))
    return jsonify({"ok": True, "persons": load_namelist()})


@app.route("/api/namelist/all", methods=["PUT"])
def replace_namelist():
    """Replace the entire list."""
    persons = request.json.get("persons", [])
    save_namelist(persons)
    return jsonify({"ok": True, "persons": load_namelist()})


@app.route("/api/docxfiles", methods=["GET"])
def list_docx_files():
    folder = request.args.get("folder", "").strip()
    if not folder:
        folder = get_article_link()
    if not folder:
        return jsonify({"files": []})
    p = Path(folder)
    if not p.exists():
        return jsonify({"files": []})
    files = sorted(f.name for f in p.iterdir() if f.suffix.lower() == ".docx")
    return jsonify({"files": files})


# ---------------------------------------------------------------------------
# Routes — output folder listing
# ---------------------------------------------------------------------------

@app.route("/api/files", methods=["GET"])
def list_files():
    article_link = get_article_link()
    if not article_link:
        return jsonify({"files": []})
    folder = Path(article_link)
    if not folder.exists():
        return jsonify({"files": []})
    files = [f.name for f in sorted(folder.iterdir()) if f.suffix in (".docx", ".pdf", ".md")]
    return jsonify({"files": files, "folder": str(folder)})


# ---------------------------------------------------------------------------
# Routes — Credit Author Statement
# ---------------------------------------------------------------------------

CREDIT_ROLES = [
    "Conceptualization", "Data Curation", "Data Processing", "Formal analysis",
    "Funding acquisition", "Investigation", "Methodology", "Original Draft",
    "Project administration", "Resources", "Review & Editing", "Software",
    "Supervision", "Validation", "Visualization",
]

CREDIT_JSON = TEMP_DIR / "CreditAuthorStatement.json"


@app.route("/api/credit", methods=["GET"])
def get_credit():
    import json
    entries = []
    if CREDIT_JSON.exists():
        try:
            entries = json.loads(CREDIT_JSON.read_text(encoding="utf-8"))
        except Exception:
            entries = []
    return jsonify({"roles": CREDIT_ROLES, "entries": entries})


@app.route("/api/credit", methods=["POST"])
def save_credit():
    import json
    entries = request.json.get("entries", [])
    TEMP_DIR.mkdir(exist_ok=True)
    CREDIT_JSON.write_text(json.dumps(entries, ensure_ascii=False, indent=2), encoding="utf-8")
    return jsonify({"ok": True})


# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Routes — Filesystem path autocomplete
# ---------------------------------------------------------------------------

@app.route("/api/ls", methods=["GET"])
def ls_dirs():
    raw = request.args.get("path", "").strip()
    if not raw:
        raw = str(Path.home())
    p = Path(raw).expanduser()
    # If path doesn't exist as a dir, treat parent as dir and last part as prefix
    if p.is_dir():
        parent, prefix = p, ""
    else:
        parent, prefix = p.parent, p.name
    if not parent.is_dir():
        return jsonify({"entries": []})
    try:
        entries = sorted(
            str(child) + ("/" if child.is_dir() else "")
            for child in parent.iterdir()
            if child.name.startswith(prefix) and not child.name.startswith(".")
        )
    except PermissionError:
        entries = []
    return jsonify({"entries": entries[:40]})


# ---------------------------------------------------------------------------
# Routes — Manuscript preview (title + abstract)
# ---------------------------------------------------------------------------

@app.route("/api/manuscript_preview", methods=["GET"])
def manuscript_preview():
    folder   = request.args.get("folder", "").strip() or get_article_link()
    filename = request.args.get("file", "").strip()
    if not folder or not filename:
        return jsonify({"title": "", "abstract": ""})
    path = Path(folder) / filename
    if not path.exists():
        return jsonify({"title": "", "abstract": ""})
    try:
        from docx import Document
        doc   = Document(str(path))
        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        title = paras[0] if paras else ""
        # Find abstract: paragraph after one containing "abstract" (case-insensitive)
        abstract = ""
        for i, text in enumerate(paras):
            if text.lower().startswith("abstract"):
                # The heading itself may contain the abstract body, or the next para does
                body = text[len("abstract"):].lstrip(": \n").strip()
                abstract = body if body else (paras[i + 1] if i + 1 < len(paras) else "")
                break
        return jsonify({"title": title, "abstract": abstract})
    except Exception as e:
        return jsonify({"title": "", "abstract": "", "error": str(e)})


# ---------------------------------------------------------------------------
# Routes — Submission Log
# ---------------------------------------------------------------------------

SUBMISSION_LOG_NAME = "submission_log.txt"


@app.route("/api/submission", methods=["GET"])
def get_submission_log():
    article_link = get_article_link()
    if not article_link:
        return jsonify({"lines": [], "journal": ""})
    log_path = Path(article_link) / SUBMISSION_LOG_NAME
    lines = []
    if log_path.exists():
        lines = [l for l in log_path.read_text(encoding="utf-8").splitlines() if l.strip()]
    # Also return current journal from .input for convenience
    data = parse_input_fields(INPUT_FILE)
    journal = data.get("journal_name", "")
    return jsonify({"lines": lines, "journal": journal})


@app.route("/api/submission", methods=["POST"])
def log_submission():
    from datetime import datetime
    article_link = get_article_link()
    if not article_link:
        return jsonify({"ok": False, "error": "article_link not set"})
    journal = request.json.get("journal", "").strip()
    if not journal:
        return jsonify({"ok": False, "error": "journal name required"})
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    line = f"{now}  →  {journal}"
    log_path = Path(article_link) / SUBMISSION_LOG_NAME
    existing = log_path.read_text(encoding="utf-8") if log_path.exists() else ""
    if existing and not existing.endswith("\n"):
        existing += "\n"
    log_path.write_text(existing + line + "\n", encoding="utf-8")
    return jsonify({"ok": True, "line": line})


# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print(f"[HappyToGo STAGE] running at http://127.0.0.1:5050")
    app.run(debug=True, port=5050)
