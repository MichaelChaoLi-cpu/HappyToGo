#!/usr/bin/env python3
"""
cover_letter_creator.py

Flow:
  1. Read .input → locate manuscript docx
  2. Call LLM → save markdown to _temp/CoverLetter.md
  3. Parse markdown → build CoverLetter.docx via python-docx
     Signature is taken from NameList (corresponding author).

Required .input keys:
  article_link   path to the manuscript folder
  file_name      manuscript filename

Optional .input keys:
  journal_name       target journal (default: [TARGET JOURNAL])
  cover_letter_name  output filename (default: CoverLetter.docx)
  corresponding      corresponding author name (must exist in NameList)
"""

from __future__ import annotations

import os
import sys
import re
import json
from datetime import date
from pathlib import Path

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR   = Path(__file__).parent
PROJECT_ROOT = SCRIPT_DIR.parent
INPUT_FILE   = PROJECT_ROOT / ".input"
TEMP_DIR     = PROJECT_ROOT / "_temp"
NAMELIST_FILE = PROJECT_ROOT / "INFOCENTER" / "NameList.json"

sys.path.insert(0, str(SCRIPT_DIR))
from llm import call_llm  # noqa: E402

# ---------------------------------------------------------------------------
# Page / font constants  (A4, 1.25" L/R, 1.0" T/B)
# ---------------------------------------------------------------------------
_EMU_INCH = 914400
_EMU_PT   = 12700

PAGE_WIDTH    = int(8.268 * _EMU_INCH)   # 7560691  ≈ A4
PAGE_HEIGHT   = int(11.693 * _EMU_INCH)  # 10692520 ≈ A4
MARGIN_LR     = int(1.25 * _EMU_INCH)    # 1143000
MARGIN_TB     = int(1.0  * _EMU_INCH)    # 914400
SIZE_BODY     = 12 * _EMU_PT             # 152400
FONT_SERIF    = "Times New Roman"
LINE_SPACING  = 1.5   # multiple


# ---------------------------------------------------------------------------
# Load .env
# ---------------------------------------------------------------------------
def _load_env():
    env_file = PROJECT_ROOT / ".env"
    if env_file.exists():
        for line in env_file.read_text().splitlines():
            line = line.strip()
            if line.startswith("export "):
                line = line[7:]
            if "=" in line and not line.startswith("#"):
                k, v = line.split("=", 1)
                os.environ.setdefault(k.strip(), v.strip().strip("'\""))


# ---------------------------------------------------------------------------
# Parse .input
# ---------------------------------------------------------------------------
def parse_input(path: Path) -> dict:
    result: dict = {}
    content = path.read_text()
    for m in re.finditer(r'(\w+)\s*=\s*[{\[](.*?)[}\]]', content, re.DOTALL):
        items = [s.strip().strip("'\"") for s in m.group(2).split(',') if s.strip().strip("'\"")]
        result[m.group(1)] = items
    for line in content.splitlines():
        m = re.match(r"""(\w+)\s*=\s*['"]([^'"]+)['"]""", line.strip())
        if m and m.group(1) not in result:
            result[m.group(1)] = m.group(2)
    return result


# ---------------------------------------------------------------------------
# NameList
# ---------------------------------------------------------------------------
def load_namelist() -> dict[str, dict]:
    if not NAMELIST_FILE.exists():
        return {}
    try:
        persons = json.loads(NAMELIST_FILE.read_text(encoding="utf-8"))
        return {
            p.get("name", "").strip(): {
                "institution": p.get("institution", "").strip(),
                "email":       p.get("email", "").strip(),
                "address":     p.get("address", "").strip(),
                "position":    p.get("position", "").strip(),
                "phone":       p.get("phone", "").strip(),
            }
            for p in persons if p.get("name", "").strip()
        }
    except Exception:
        return {}


def build_signature_lines(corresponding: str, namelist: dict[str, dict]) -> list[str]:
    info = namelist.get(corresponding, {})
    lines = [corresponding]
    if info.get("position"):
        lines.append(info["position"])
    if info.get("institution"):
        lines.append(info["institution"])
    if info.get("address"):
        lines.append(info["address"])
    if info.get("phone"):
        lines.append(f"Tel: {info['phone']}")
    if info.get("email"):
        lines.append(f"Email: {info['email']}")
    return lines


# ---------------------------------------------------------------------------
# Extract from manuscript docx
# ---------------------------------------------------------------------------
def extract_docx_text(docx_path: Path, max_chars: int = 6000) -> str:
    from docx import Document
    doc = Document(str(docx_path))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())[:max_chars]


def extract_docx_title(docx_path: Path, fallback: str) -> str:
    from docx import Document
    doc = Document(str(docx_path))
    return next((p.text.strip() for p in doc.paragraphs if p.text.strip()), fallback)


# ---------------------------------------------------------------------------
# LLM → Markdown
# ---------------------------------------------------------------------------
def generate_markdown(manuscript_text: str, manuscript_title: str,
                      journal_name: str, today_str: str,
                      signature_lines: list[str]) -> str:
    sig_block = "\n\n".join(signature_lines)
    prompt = f"""You are an expert academic editor. Write a complete cover letter for a journal submission in Markdown format.

MANUSCRIPT TITLE:
{manuscript_title}

TARGET JOURNAL:
{journal_name}

DATE:
{today_str}

MANUSCRIPT CONTENT (excerpt):
{manuscript_text}

OUTPUT FORMAT (strict Markdown, reproduce this structure exactly, replacing only the bracketed placeholders):

---
{manuscript_title}

{journal_name}

{today_str}

Dear Editor,

[Paragraph 1: Introduce the manuscript, its scope, data, and methods. Plain prose, no bullet points.]

[Paragraph 2: Summarize 2-3 key findings. Plain prose, no bullet points.]

[Paragraph 3: Explain the methodological and substantive contributions.]

[Paragraph 4: Explain why it fits the journal and invite the editor to reach out.]

Thank you for considering our manuscript. We look forward to your feedback and are eager to contribute to advancing research in this field.

Sincerely,

{sig_block}
---

RULES:
- Keep the first three lines exactly as shown: manuscript title, journal name, date.
- Replace the bracketed placeholders with real content based on the manuscript.
- Write in formal academic prose. No bullet points. No numbered lists.
- Do NOT add any headings (no # ## ###).
- Do NOT use bold (**text**) or italic (*text*) anywhere.
- Keep each paragraph as a single unbroken block of text.
- Output ONLY the content between the --- markers (do not include the --- lines).
"""
    return call_llm(prompt).strip()


def strip_markdown_emphasis(text: str) -> str:
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'__(.+?)__',     r'\1', text, flags=re.DOTALL)
    text = re.sub(r'\*(.+?)\*',     r'\1', text, flags=re.DOTALL)
    text = re.sub(r'(?<!\w)_(.+?)_(?!\w)', r'\1', text, flags=re.DOTALL)
    return text


# ---------------------------------------------------------------------------
# Parse markdown → structured sections
# ---------------------------------------------------------------------------
def parse_markdown(md: str) -> dict:
    """
    Returns:
      title, journal, date_str,
      body_paras (list[str]),   ← includes "Thank you..." paragraph
    """
    blocks = [b.strip() for b in re.split(r'\n{2,}', md.strip()) if b.strip()]

    title     = blocks[0] if len(blocks) > 0 else ""
    journal   = blocks[1] if len(blocks) > 1 else ""
    date_str  = blocks[2] if len(blocks) > 2 else ""

    # Everything after "Dear Editor," up to (not including) "Sincerely,"
    body_paras: list[str] = []
    in_body = False
    for block in blocks[3:]:
        low = block.lower().strip()
        if low.startswith("dear editor"):
            in_body = True
            continue
        if low.startswith("sincerely"):
            break
        if in_body:
            body_paras.append(block)

    return {"title": title, "journal": journal, "date": date_str, "body": body_paras}


# ---------------------------------------------------------------------------
# Build CoverLetter.docx from scratch
# ---------------------------------------------------------------------------
def build_docx(parsed: dict, signature_lines: list[str], output_path: Path):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # Page layout (A4)
    sec = doc.sections[0]
    sec.page_width    = PAGE_WIDTH
    sec.page_height   = PAGE_HEIGHT
    sec.left_margin   = MARGIN_LR
    sec.right_margin  = MARGIN_LR
    sec.top_margin    = MARGIN_TB
    sec.bottom_margin = MARGIN_TB

    # Remove default empty paragraph
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    def _spacing(para, align=None, space_before=0, space_after=0):
        fmt = para.paragraph_format
        fmt.space_before = space_before
        fmt.space_after  = space_after
        pPr = para._element.get_or_add_pPr()
        sp  = pPr.get_or_add_spacing()
        sp.set(qn("w:lineRule"), "auto")
        sp.set(qn("w:line"), str(int(LINE_SPACING * 240)))
        if align is not None:
            para.alignment = align

    def _run(para, text, bold=False):
        r = para.add_run(text)
        r.bold       = bold
        r.font.size  = SIZE_BODY
        r.font.name  = FONT_SERIF
        return r

    def _add(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, space_before=0, space_after=0):
        p = doc.add_paragraph()
        _spacing(p, align=align, space_before=space_before, space_after=space_after)
        _run(p, text, bold=bold)
        return p

    def _blank():
        p = doc.add_paragraph()
        _spacing(p)
        _run(p, "")

    # ── Header block ──────────────────────────────────────────────────────
    _add(parsed["title"],   align=WD_ALIGN_PARAGRAPH.LEFT, bold=True)
    _add(parsed["journal"], align=WD_ALIGN_PARAGRAPH.LEFT, bold=True)
    _add(parsed["date"],    align=WD_ALIGN_PARAGRAPH.RIGHT)

    _blank()
    _add("Dear Editor,", align=WD_ALIGN_PARAGRAPH.LEFT)
    _blank()

    # ── Body paragraphs (JUSTIFY) ─────────────────────────────────────────
    for i, para_text in enumerate(parsed["body"]):
        _add(para_text, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        if i < len(parsed["body"]) - 1:
            _blank()

    # ── Closing ───────────────────────────────────────────────────────────
    _blank()
    _add("Sincerely,", align=WD_ALIGN_PARAGRAPH.LEFT)
    _blank()

    # ── Signature from NameList ───────────────────────────────────────────
    for line in signature_lines:
        _add(line, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.save(str(output_path))
    print(f"[✓] Cover letter saved: {output_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    _load_env()

    build_only = "--build" in sys.argv

    if not INPUT_FILE.exists():
        sys.exit(f"[Error] .input not found: {INPUT_FILE}")

    cfg               = parse_input(INPUT_FILE)
    article_link      = (cfg.get("article_link") or "").strip()
    file_name         = (cfg.get("file_name") or "").strip()
    journal_name      = (cfg.get("journal_name") or "[TARGET JOURNAL]").strip()
    cover_letter_name = (cfg.get("cover_letter_name") or "CoverLetter.docx").strip()
    corresponding     = (cfg.get("corresponding") or "").strip()

    if not article_link or not file_name:
        sys.exit("[Error] .input must define article_link and file_name")

    manuscript_path = Path(article_link) / file_name
    if not manuscript_path.exists():
        sys.exit(f"[Error] Manuscript not found: {manuscript_path}")

    journal_folder = Path(article_link) / journal_name
    journal_folder.mkdir(exist_ok=True)
    output_path = journal_folder / cover_letter_name
    TEMP_DIR.mkdir(exist_ok=True)
    temp_md = TEMP_DIR / "CoverLetter.md"

    # Build signature from NameList
    namelist = load_namelist()
    signature_lines = build_signature_lines(corresponding, namelist)

    if build_only:
        if not temp_md.exists():
            sys.exit(f"[Error] No markdown found at {temp_md}. Run without --build first.")
        print(f"[→] Using existing markdown: {temp_md}")
    else:
        manuscript_text  = extract_docx_text(manuscript_path)
        manuscript_title = extract_docx_title(manuscript_path, file_name)
        print(f"[→] Title:   {manuscript_title[:80]}...")
        print(f"[→] Journal: {journal_name}")

        d = date.today()
        suffix = "th" if 11 <= d.day <= 13 else {1:"st",2:"nd",3:"rd"}.get(d.day % 10, "th")
        today_str = d.strftime(f"%B {d.day}{suffix}, %Y")

        print("[→] Calling LLM to generate markdown...")
        md_content = generate_markdown(
            manuscript_text, manuscript_title, journal_name, today_str, signature_lines
        )
        md_content = strip_markdown_emphasis(md_content)
        temp_md.write_text(md_content, encoding="utf-8")
        print(f"[→] Markdown saved: {temp_md}")

    md_text = temp_md.read_text(encoding="utf-8")
    parsed  = parse_markdown(md_text)
    print(f"[→] Parsed: {len(parsed['body'])} body paragraphs")

    build_docx(parsed, signature_lines, output_path)


if __name__ == "__main__":
    main()
