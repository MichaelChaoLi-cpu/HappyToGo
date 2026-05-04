#!/usr/bin/env python3
"""
title_creator.py

Usage:
  python title_creator.py          → write _temp/Title.md then build Title.docx
  python title_creator.py --build  → build Title.docx from existing _temp/Title.md

_temp/Title.md format:
  # Title
  <manuscript title>

  # Authors
  Name[1], Name[1], CorrespondingName*[1]

  # Affiliations
  [1] Urban Institute & School of Engineering, Kyushu University, Japan

  # Corresponding
  Name | email@address | Full postal address

Required .input keys: article_link, file_name
Optional .input keys: title_name, authors, corresponding
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

SCRIPT_DIR     = Path(__file__).parent
PROJECT_ROOT   = SCRIPT_DIR.parent
INPUT_FILE     = PROJECT_ROOT / ".input"
TEMP_DIR       = PROJECT_ROOT / "_temp"
TEMP_MD        = TEMP_DIR / "Title.md"
NAMELIST_FILE  = PROJECT_ROOT / "INFOCENTER" / "NameList.json"

# ---------------------------------------------------------------------------
# Page / font constants (US Letter, 1.25" L/R, 1.0" T/B)
# ---------------------------------------------------------------------------
_EMU_INCH  = 914400
_EMU_PT    = 12700

PAGE_WIDTH       = int(8.5  * _EMU_INCH)   # 7772400
PAGE_HEIGHT      = int(11.0 * _EMU_INCH)   # 10058400
MARGIN_LR        = int(1.25 * _EMU_INCH)   # 1143000
MARGIN_TB        = int(1.0  * _EMU_INCH)   # 914400

SIZE_TITLE       = 16 * _EMU_PT   # 203200
SIZE_BODY        = 12 * _EMU_PT   # 152400
SIZE_NORMAL      = int(10.5 * _EMU_PT)  # 133350
FONT_SERIF       = "Times New Roman"

SPACE_BEFORE_6PT = int(6 * _EMU_PT)   # 76200


# ---------------------------------------------------------------------------
# NameList helpers
# ---------------------------------------------------------------------------
def load_namelist() -> dict[str, dict]:
    """Return {name: {institution, email, address}} from NameList.json."""
    import json
    if not NAMELIST_FILE.exists():
        return {}
    try:
        persons = json.loads(NAMELIST_FILE.read_text(encoding="utf-8"))
        return {
            p.get("name", "").strip(): {
                "institution": p.get("institution", "").strip(),
                "email":       p.get("email", "").strip(),
                "address":     p.get("address", "").strip(),
                "position":    p.get("position", "").strip(),
                "phone":       p.get("phone", "").strip(),
            }
            for p in persons if p.get("name", "").strip()
        }
    except Exception:
        return {}


def derive_affiliations(authors: list[str], namelist: dict[str, dict]) -> tuple[list[str], dict[str, str]]:
    """
    Build affiliation list and per-author affiliation-number map.
    Authors sharing the same institution get the same [N].
    Returns (["[1] Inst A", "[2] Inst B", ...], {author_name: "1", ...})
    """
    inst_to_num: dict[str, str] = {}
    aff_map: dict[str, str] = {}
    counter = 1

    for name in authors:
        inst = namelist.get(name, {}).get("institution", "")
        if inst and inst not in inst_to_num:
            inst_to_num[inst] = str(counter)
            counter += 1
        aff_map[name] = inst_to_num.get(inst, "1") if inst else "1"

    affiliations = [f"[{num}] {inst}" for inst, num in
                    sorted(inst_to_num.items(), key=lambda x: int(x[1]))]
    if not affiliations:
        affiliations = ["[1] "]
    return affiliations, aff_map


# ---------------------------------------------------------------------------
# Parse .input
# ---------------------------------------------------------------------------
def parse_input(path: Path) -> dict:
    content = path.read_text()
    result: dict = {}
    for m in re.finditer(r'(\w+)\s*=\s*[{\[](.*?)[}\]]', content, re.DOTALL):
        key   = m.group(1)
        items = [s.strip().strip("'\"") for s in m.group(2).split(',') if s.strip().strip("'\"")]
        result[key] = items
    for line in content.splitlines():
        m = re.match(r"""(\w+)\s*=\s*['"]([^'"]+)['"]""", line.strip())
        if m and m.group(1) not in result:
            result[m.group(1)] = m.group(2)
    return result


# ---------------------------------------------------------------------------
# Extract manuscript title (first non-empty paragraph)
# ---------------------------------------------------------------------------
def extract_title(docx_path: Path) -> str:
    from docx import Document
    doc = Document(str(docx_path))
    return next((p.text.strip() for p in doc.paragraphs if p.text.strip()), "")


# ---------------------------------------------------------------------------
# Write _temp/Title.md
# ---------------------------------------------------------------------------
def write_temp_md(title: str, authors: list[str], corresponding: str,
                  affiliations: list[str], corresponding_text: str,
                  aff_map: dict[str, str] | None = None):
    TEMP_DIR.mkdir(exist_ok=True)

    def _aff_num(name: str) -> str:
        return aff_map[name] if aff_map and name in aff_map else "1"

    author_line = ", ".join(
        f"{a}*[{_aff_num(a)}]" if a == corresponding else f"{a}[{_aff_num(a)}]"
        for a in authors
    )
    aff_block = "\n".join(affiliations) if affiliations else "[1] "

    md = (
        f"# Title\n{title}\n\n"
        f"# Authors\n{author_line}\n\n"
        f"# Affiliations\n{aff_block}\n\n"
        f"# Corresponding\n{corresponding_text}\n"
    )
    TEMP_MD.write_text(md, encoding="utf-8")
    print(f"[→] Markdown saved: {TEMP_MD}")


# ---------------------------------------------------------------------------
# Parse _temp/Title.md
# ---------------------------------------------------------------------------
def read_temp_md() -> dict:
    text = TEMP_MD.read_text(encoding="utf-8")

    def get_section(name: str) -> str:
        m = re.search(rf"^# {name}\s*\n(.*?)(?=^# |\Z)", text, re.MULTILINE | re.DOTALL)
        return m.group(1).strip() if m else ""

    title        = get_section("Title")
    authors_raw  = get_section("Authors")
    aff_raw      = get_section("Affiliations")
    corr_raw     = get_section("Corresponding")

    authors      = []
    corresponding = ""
    aff_map      = {}

    for token in re.split(r",\s*", authors_raw):
        token = token.strip()
        if not token:
            continue
        m = re.search(r"\[(\d+)\]", token)
        aff_num = m.group(1) if m else "1"
        name = re.sub(r"\[\d+\]", "", token).rstrip("*").strip()
        is_corr = "*" in token.replace(f"[{aff_num}]", "")
        authors.append(name)
        aff_map[name] = aff_num
        if is_corr:
            corresponding = name

    if not corresponding and authors:
        corresponding = authors[-1]

    affiliations = [line.strip() for line in aff_raw.splitlines() if line.strip()]

    corr_parts   = [p.strip() for p in corr_raw.split("|")]
    corr_name    = corr_parts[0] if len(corr_parts) > 0 else corresponding
    corr_email   = corr_parts[1] if len(corr_parts) > 1 else ""
    corr_address = corr_parts[2] if len(corr_parts) > 2 else ""

    return {
        "title": title,
        "authors": authors,
        "corresponding": corresponding,
        "aff_map": aff_map,
        "affiliations": affiliations,
        "corr_name": corr_name,
        "corr_email": corr_email,
        "corr_address": corr_address,
    }


# ---------------------------------------------------------------------------
# Build Title.docx from scratch
# ---------------------------------------------------------------------------
def build_docx(data: dict, output_path: Path):
    from docx import Document
    from docx.shared import Pt  # noqa: F401 — kept for potential future use
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page layout
    sec = doc.sections[0]
    sec.page_width   = PAGE_WIDTH
    sec.page_height  = PAGE_HEIGHT
    sec.left_margin  = MARGIN_LR
    sec.right_margin = MARGIN_LR
    sec.top_margin   = MARGIN_TB
    sec.bottom_margin = MARGIN_TB

    # Remove default empty paragraph that Document() creates
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    def _set_spacing(para, space_before=0, space_after=0, line_spacing=None):
        fmt = para.paragraph_format
        fmt.space_before = space_before
        fmt.space_after  = space_after
        if line_spacing is not None:
            # Set multiple line spacing via OOXML directly (lineRule="auto", line=240*factor)
            pPr = para._element.get_or_add_pPr()
            spacing = pPr.get_or_add_spacing()
            spacing.set(qn("w:lineRule"), "auto")
            spacing.set(qn("w:line"), str(int(line_spacing * 240)))

    def _add_run(para, text, bold=None, superscript=False,
                 size=SIZE_BODY, font_name=FONT_SERIF):
        r = para.add_run(text)
        r.bold = bold
        r.font.superscript = superscript
        r.font.size  = size
        r.font.name  = font_name
        return r

    # ── Para 0: Title ──────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p_title, space_before=0, space_after=0, line_spacing=1.5)
    _add_run(p_title, data["title"], bold=True, size=SIZE_TITLE)

    # ── Para 1: "Authors" label ────────────────────────────────────────────
    p_authors_label = doc.add_paragraph()
    p_authors_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(p_authors_label, space_before=SPACE_BEFORE_6PT, space_after=0)
    _add_run(p_authors_label, "Authors", bold=True)

    # ── Para 2: Author names with superscript affiliation numbers ──────────
    p_authors = doc.add_paragraph()
    p_authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p_authors, space_before=0, space_after=0)

    for i, name in enumerate(data["authors"]):
        display = f"{name}*" if name == data["corresponding"] else name
        _add_run(p_authors, display)
        _add_run(p_authors, data["aff_map"].get(name, "1"), superscript=True)
        if i < len(data["authors"]) - 1:
            _add_run(p_authors, ", ")

    # ── Para 3: Empty separator ────────────────────────────────────────────
    p_sep = doc.add_paragraph()
    _set_spacing(p_sep, space_before=0, space_after=0)

    # ── Para 4: "Affiliations" label ───────────────────────────────────────
    p_aff_label = doc.add_paragraph()
    p_aff_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(p_aff_label, space_before=0, space_after=0)
    _add_run(p_aff_label, "Affiliations", bold=True)

    # ── Para 5+: Affiliation lines ─────────────────────────────────────────
    for aff_line in data["affiliations"]:
        m = re.match(r"\[(\d+)\]\s*(.*)", aff_line)
        num  = m.group(1) if m else "1"
        desc = m.group(2) if m else aff_line

        p_aff = doc.add_paragraph()
        p_aff.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_spacing(p_aff, space_before=0, space_after=0)
        _add_run(p_aff, num, superscript=True)
        _add_run(p_aff, f" {desc}")

    # ── Para N: Corresponding line ─────────────────────────────────────────
    p_corr = doc.add_paragraph()
    p_corr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(p_corr, space_before=0, space_after=0)

    _add_run(p_corr, "* ")
    _add_run(p_corr, f"Correspondent to: {data['corr_name']}, ")

    # Hyperlink for email
    if data["corr_email"]:
        url   = f"mailto:{data['corr_email']}"
        r_id  = p_corr.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hl = OxmlElement("w:hyperlink")
        hl.set(qn("r:id"), r_id)

        hl_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), FONT_SERIF)
        rFonts.set(qn("w:hAnsi"), FONT_SERIF)
        rPr.append(rFonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(SIZE_BODY // (_EMU_PT // 2)))  # half-points
        rPr.append(sz)
        hl_run.append(rPr)

        t_elem = OxmlElement("w:t")
        t_elem.text = data["corr_email"]
        hl_run.append(t_elem)
        hl.append(hl_run)
        p_corr._element.append(hl)

    if data["corr_address"]:
        _add_run(p_corr, f", {data['corr_address']}")

    doc.save(str(output_path))
    print(f"[✓] Title page saved: {output_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    build_only = "--build" in sys.argv

    if not INPUT_FILE.exists():
        sys.exit(f"[Error] .input not found: {INPUT_FILE}")

    cfg          = parse_input(INPUT_FILE)
    article_link = (cfg.get("article_link") or "").strip()
    file_name    = (cfg.get("file_name") or "").strip()
    title_name   = (cfg.get("title_name") or "Title.docx").strip()

    if not article_link or not file_name:
        sys.exit("[Error] .input must define article_link and file_name")

    manuscript_path = Path(article_link) / file_name
    if not manuscript_path.exists():
        sys.exit(f"[Error] Manuscript not found: {manuscript_path}")

    output_path = Path(article_link) / title_name

    if build_only:
        if not TEMP_MD.exists():
            sys.exit(f"[Error] No markdown found at {TEMP_MD}. Run without --build first.")
        print(f"[→] Using existing markdown: {TEMP_MD}")
    else:
        title         = extract_title(manuscript_path)
        authors: list[str] = cfg.get("authors") or []
        corresponding: str = (cfg.get("corresponding") or "").strip()

        if not authors:
            authors       = ["Author"]
            corresponding = "Author"
        elif not corresponding:
            corresponding = authors[-1]

        # Derive affiliations and contact info from NameList
        namelist = load_namelist()
        affiliations, aff_map = derive_affiliations(authors, namelist)

        corr_info    = namelist.get(corresponding, {})
        corr_email   = corr_info.get("email", "")
        corr_address = corr_info.get("address", "")
        corr_text    = f"{corresponding} | {corr_email} | {corr_address}"

        write_temp_md(title, authors, corresponding, affiliations, corr_text, aff_map)

    data = read_temp_md()
    print(f"[→] Title:        {data['title'][:70]}...")
    print(f"[→] Authors:      {data['authors']}")
    print(f"[→] Corresponding:{data['corresponding']}")

    build_docx(data, output_path)


if __name__ == "__main__":
    main()
