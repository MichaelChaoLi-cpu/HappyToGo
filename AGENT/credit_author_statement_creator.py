#!/usr/bin/env python3
"""
credit_author_statement_creator.py

Usage:
  python credit_author_statement_creator.py         → build docx from _temp/CreditAuthorStatement.json
  python credit_author_statement_creator.py --build → same (alias, always reads json)

_temp/CreditAuthorStatement.json format:
  [
    {"name": "Li Chao",        "roles": ["Conceptualization", "Methodology", ...]},
    {"name": "Mi Jie",         "roles": ["Review & Editing"]},
    ...
  ]

Required .input keys: article_link
Optional .input keys: credit_name (default: CreditAuthorStatement.docx)
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

SCRIPT_DIR   = Path(__file__).parent
PROJECT_ROOT = SCRIPT_DIR.parent
INPUT_FILE   = PROJECT_ROOT / ".input"
TEMP_DIR     = PROJECT_ROOT / "_temp"
TEMP_JSON    = TEMP_DIR / "CreditAuthorStatement.json"

# Page / font constants (US Letter, 1.25" L/R, 1.0" T/B, Times New Roman 12pt)
_EMU_INCH     = 914400
_EMU_PT       = 12700
PAGE_WIDTH    = int(8.5  * _EMU_INCH)
PAGE_HEIGHT   = int(11.0 * _EMU_INCH)
MARGIN_LR     = int(1.25 * _EMU_INCH)
MARGIN_TB     = int(1.0  * _EMU_INCH)
SIZE_BODY     = 12 * _EMU_PT
FONT_SERIF    = "Times New Roman"
SPACE_BEFORE  = int(6 * _EMU_PT)   # 76200 — matches Acknowledgement style

# Canonical role list (order preserved from template)
CREDIT_ROLES = [
    "Conceptualization",
    "Methodology",
    "Software",
    "Validation",
    "Formal analysis",
    "Investigation",
    "Data Curation",
    "Data Processing",
    "Original Draft",
    "Review & Editing",
    "Visualization",
    "Resources",
    "Supervision",
    "Project administration",
    "Funding acquisition",
]


# ---------------------------------------------------------------------------
# Parse .input
# ---------------------------------------------------------------------------
def parse_input(path: Path) -> dict:
    result: dict = {}
    content = path.read_text()
    for m in re.finditer(r'(\w+)\s*=\s*[{\[](.*?)[}\]]', content, re.DOTALL):
        items = [s.strip().strip("'\"") for s in m.group(2).split(',') if s.strip().strip("'\"")]
        result[m.group(1)] = items
    for line in content.splitlines():
        m = re.match(r"""(\w+)\s*=\s*['"]([^'"]+)['"]""", line.strip())
        if m and m.group(1) not in result:
            result[m.group(1)] = m.group(2)
    return result


# ---------------------------------------------------------------------------
# Build CreditAuthorStatement.docx from scratch
# ---------------------------------------------------------------------------
def build_docx(entries: list[dict], output_path: Path):
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document()

    sec = doc.sections[0]
    sec.page_width    = PAGE_WIDTH
    sec.page_height   = PAGE_HEIGHT
    sec.left_margin   = MARGIN_LR
    sec.right_margin  = MARGIN_LR
    sec.top_margin    = MARGIN_TB
    sec.bottom_margin = MARGIN_TB

    # Remove default empty paragraph
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    def _add(space_before=SPACE_BEFORE, space_after=0):
        p = doc.add_paragraph()
        fmt = p.paragraph_format
        fmt.space_before = space_before
        fmt.space_after  = space_after
        # Single line spacing via XML
        pPr = p._element.get_or_add_pPr()
        sp  = pPr.get_or_add_spacing()
        sp.set(qn("w:lineRule"), "auto")
        sp.set(qn("w:line"), "240")   # 1.0× (240 units = single)
        return p

    def _run(para, text, bold=False):
        r = para.add_run(text)
        r.bold      = bold
        r.font.size = SIZE_BODY
        r.font.name = FONT_SERIF
        return r

    # Title: "Credit Author Statement"
    p0 = _add(space_before=0)
    _run(p0, "Credit Author Statement", bold=True)

    # Empty separator
    _add(space_before=0)

    # Author lines: bold name + normal roles
    for entry in entries:
        name  = entry.get("name", "").strip()
        roles = entry.get("roles", [])
        if not name or not roles:
            continue
        p = _add()
        _run(p, name, bold=True)
        _run(p, f": {', '.join(roles)}")

    doc.save(str(output_path))
    print(f"[✓] CreditAuthorStatement saved: {output_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    if not INPUT_FILE.exists():
        sys.exit(f"[Error] .input not found: {INPUT_FILE}")
    if not TEMP_JSON.exists():
        sys.exit(f"[Error] No JSON at {TEMP_JSON}. Fill in roles via the web UI first.")

    cfg          = parse_input(INPUT_FILE)
    article_link = cfg.get("article_link", "").strip()
    credit_name  = cfg.get("credit_name", "CreditAuthorStatement.docx").strip()

    if not article_link:
        sys.exit("[Error] .input must define article_link")

    entries = json.loads(TEMP_JSON.read_text(encoding="utf-8"))
    for e in entries:
        print(f"  {e['name']}: {', '.join(e['roles'])}")

    output_path = Path(article_link) / credit_name
    build_docx(entries, output_path)


if __name__ == "__main__":
    main()
