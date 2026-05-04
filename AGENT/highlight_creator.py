#!/usr/bin/env python3
"""
highlight_creator.py

Usage:
  python highlight_creator.py          → call Gemini, write _temp/Highlights.md, build docx
  python highlight_creator.py --build  → build docx from existing _temp/Highlights.md

_temp/Highlights.md format (one highlight per line, edit freely):
  This study uses global survey data across 22 countries to examine loneliness.
  Rural residents report consistently higher loneliness than urban counterparts.
  ...

Required .input keys: article_link, file_name
Optional .input keys: highlight_name (default: Highlights.docx)
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

SCRIPT_DIR   = Path(__file__).parent
PROJECT_ROOT = SCRIPT_DIR.parent
INPUT_FILE   = PROJECT_ROOT / ".input"
TEMP_DIR     = PROJECT_ROOT / "_temp"
TEMP_MD      = TEMP_DIR / "Highlights.md"

# Page / font constants (US Letter, 1.25" L/R, 1.0" T/B, Times New Roman 12pt)
_EMU_INCH  = 914400
_EMU_PT    = 12700
PAGE_WIDTH  = int(8.5  * _EMU_INCH)
PAGE_HEIGHT = int(11.0 * _EMU_INCH)
MARGIN_LR   = int(1.25 * _EMU_INCH)
MARGIN_TB   = int(1.0  * _EMU_INCH)
SIZE_BODY   = 12 * _EMU_PT
FONT_SERIF  = "Times New Roman"

sys.path.insert(0, str(SCRIPT_DIR))
from llm import call_llm


# ---------------------------------------------------------------------------
# Load .env
# ---------------------------------------------------------------------------
def _load_env():
    import os
    env_file = PROJECT_ROOT / ".env"
    if env_file.exists():
        for line in env_file.read_text().splitlines():
            line = line.strip()
            if line.startswith("export "):
                line = line[7:]
            if "=" in line and not line.startswith("#"):
                k, v = line.split("=", 1)
                os.environ.setdefault(k.strip(), v.strip().strip("'\""))


# ---------------------------------------------------------------------------
# Parse .input
# ---------------------------------------------------------------------------
def parse_input(path: Path) -> dict:
    result: dict = {}
    for line in path.read_text().splitlines():
        m = re.match(r"""(\w+)\s*=\s*['"]([^'"]+)['"]""", line.strip())
        if m:
            result[m.group(1)] = m.group(2)
    return result


# ---------------------------------------------------------------------------
# Extract manuscript text
# ---------------------------------------------------------------------------
def extract_docx_text(docx_path: Path, max_chars: int = 8000) -> str:
    from docx import Document
    doc = Document(str(docx_path))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())[:max_chars]


# ---------------------------------------------------------------------------
# Call Gemini
# ---------------------------------------------------------------------------
def generate_highlights(manuscript_text: str) -> list[str]:
    prompt = f"""You are an expert academic editor. Write exactly 5 highlights for a journal article.

MANUSCRIPT CONTENT (excerpt):
{manuscript_text}

STRICT REQUIREMENTS:
- Exactly 5 highlights, one per line.
- Each highlight MUST be a complete sentence (subject + verb + object/complement).
- Each highlight MUST be 85 characters or fewer INCLUDING spaces. Count carefully before outputting.
- Plain text only — no bullet points, no dashes, no numbers, no markdown, no bold.
- Do NOT use parentheses ( ) for any explanations or clarifications.
- Capture the most important findings and contributions: data/methods, key results, implications.
- Output ONLY the 5 highlights, one per line, nothing else.

EXAMPLE FORMAT (style reference only):
This study uses global survey data across 22 countries to examine rural-urban loneliness.
Rural residents report consistently higher loneliness than their urban counterparts.
Social support partially mediates the rural-urban disparity in loneliness levels.
Cross-national heterogeneity suggests the loneliness gap is highly context-dependent.
Strengthening social support may reduce loneliness in disadvantaged residential settings.
"""
    raw = call_llm(prompt)
    lines = [l.strip() for l in raw.strip().splitlines() if l.strip()]
    # Strip any leading bullets/numbers Gemini might add despite instructions
    lines = [re.sub(r'^[\-\*\d\.]+\s*', '', l) for l in lines]
    return lines[:5]


# ---------------------------------------------------------------------------
# Write _temp/Highlights.md
# ---------------------------------------------------------------------------
def write_temp_md(highlights: list[str]):
    TEMP_DIR.mkdir(exist_ok=True)
    TEMP_MD.write_text("\n".join(highlights) + "\n", encoding="utf-8")
    print(f"[→] Markdown saved: {TEMP_MD}")


# ---------------------------------------------------------------------------
# Read and validate _temp/Highlights.md
# ---------------------------------------------------------------------------
def read_temp_md() -> list[str]:
    lines = [l.strip() for l in TEMP_MD.read_text(encoding="utf-8").splitlines() if l.strip()]
    if not lines:
        sys.exit(f"[Error] {TEMP_MD} is empty.")
    for i, line in enumerate(lines, 1):
        if len(line) > 85:
            print(f"[!] Line {i} is {len(line)} chars (exceeds 85): {line!r}")
    return lines


# ---------------------------------------------------------------------------
# Write Highlights.docx — built from scratch
# ---------------------------------------------------------------------------
def write_highlights_docx(output_path: Path, highlights: list[str]):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    sec = doc.sections[0]
    sec.page_width    = PAGE_WIDTH
    sec.page_height   = PAGE_HEIGHT
    sec.left_margin   = MARGIN_LR
    sec.right_margin  = MARGIN_LR
    sec.top_margin    = MARGIN_TB
    sec.bottom_margin = MARGIN_TB

    # Remove default empty paragraph
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    def _add(text, align=None, bold=False):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        r = p.add_run(text)
        r.bold      = bold
        r.font.size = SIZE_BODY
        r.font.name = FONT_SERIF

    # "Highlights:" label (bold, no explicit alignment → inherits Normal)
    _add("Highlights:", bold=True)

    # 5 highlight sentences (JUSTIFY)
    for h in highlights:
        _add(h, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    doc.save(str(output_path))
    print(f"[✓] Highlights saved: {output_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    _load_env()
    build_only = "--build" in sys.argv

    if not INPUT_FILE.exists():
        sys.exit(f"[Error] .input not found: {INPUT_FILE}")

    cfg            = parse_input(INPUT_FILE)
    article_link   = cfg.get("article_link", "").strip()
    file_name      = cfg.get("file_name", "").strip()
    highlight_name = cfg.get("highlight_name", "Highlights.docx").strip()

    if not article_link or not file_name:
        sys.exit("[Error] .input must define article_link and file_name")

    manuscript_path = Path(article_link) / file_name
    if not manuscript_path.exists():
        sys.exit(f"[Error] Manuscript not found: {manuscript_path}")

    output_path = Path(article_link) / highlight_name

    if build_only:
        if not TEMP_MD.exists():
            sys.exit(f"[Error] No markdown at {TEMP_MD}. Run without --build first.")
        print(f"[→] Using existing markdown: {TEMP_MD}")
    else:
        print(f"[→] Reading manuscript: {manuscript_path}")
        manuscript_text = extract_docx_text(manuscript_path)
        print("[→] Calling Gemini to generate highlights...")
        highlights = generate_highlights(manuscript_text)
        write_temp_md(highlights)

    highlights = read_temp_md()
    for i, h in enumerate(highlights, 1):
        print(f"  [{i}] ({len(h):2d} chars) {h}")

    write_highlights_docx(output_path, highlights)


if __name__ == "__main__":
    main()
